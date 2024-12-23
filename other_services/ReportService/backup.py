import os
import json
from flask import Flask, request, send_file
from docx import Document
from io import BytesIO
from flask_cors import CORS
from docx.shared import Pt, Cm
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement

app = Flask(__name__)
CORS(app)

# Đường dẫn template JSON
TEMPLATE_PATH = "template/bao_cao_huan_luyen_thang.json"

# Hàm xử lý nội dung ô
def clean_cell_content(content):
    if content is None or content.strip() == "":
        return ""  # Trả về chuỗi rỗng nếu nội dung trống
    return content.strip()

# Hàm xóa triệt để nội dung trong ô
def clear_cell_content(cell):
    tc = cell._tc  # Truy cập XML của ô
    # Xóa tất cả các đoạn văn bản
    for paragraph in cell.paragraphs:
        p = paragraph._element
        if p in tc:
            tc.remove(p)

    # Đảm bảo ô không chứa đoạn trống
    cell.text = "\t"

# Hàm kiểm tra và đồng nhất số lượng cột trong các hàng
def normalize_rows(rows):
    max_columns = max(len(row) for row in rows)  # Tìm số cột lớn nhất
    for row in rows:
        while len(row) < max_columns:  # Bổ sung các ô trống nếu thiếu
            row.append("")
    return rows

# Hàm tạo file Word
def create_word_file(template, data):
    doc = Document()

    # Xoay trang ngang
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = section.page_height, section.page_width

    # Xử lý từng phần tử trong template
    for item in template["template"]:
        if item["type"] == "title":
            p = doc.add_paragraph(item["content"])
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.runs[0]
            run.bold = True
            run.font.name = "Times New Roman"
            run.font.size = Pt(14)
        elif item["type"] == "subtitle":
            p = doc.add_paragraph(item["content"])
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.runs[0]
            run.font.name = "Times New Roman"
            run.font.size = Pt(12)
        elif item["type"] == "heading":
            p = doc.add_paragraph(item["content"])
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = p.runs[0]
            run.bold = True
            run.font.name = "Times New Roman"
            run.font.size = Pt(12)
        elif item["type"] == "text":
            p = doc.add_paragraph(item["content"])
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = p.runs[0]
            run.bold = False
            run.font.name = "Times New Roman"
            run.font.size = Pt(11)
        elif item["type"] == "table":
            table_id = item["id"]
            table_data = item["data"]
            total_width_percentage = sum(table_data["column_widths"])
            column_widths_cm = [Cm(24 * w / total_width_percentage) for w in table_data["column_widths"]]

            # Lấy dữ liệu bảng từ template và client
            rows = normalize_rows(table_data["rows"] + data.get(table_id, {}).get("rows", []))

            # Tạo bảng
            table = doc.add_table(rows=len(rows), cols=len(rows[0]), style="Table Grid")

            # Cài đặt chiều rộng cột
            for col_idx, col_width in enumerate(column_widths_cm):
                for cell in table.columns[col_idx].cells:
                    cell.width = col_width

            # Điền dữ liệu vào bảng
            for row_idx, row in enumerate(rows):
                row_cells = table.rows[row_idx].cells
                for col_idx, cell_data in enumerate(row):
                    cleaned_content = clean_cell_content(cell_data)
                    row_cells[col_idx].text = cleaned_content

                    # In đậm dòng tiêu đề trong template
                    if row_idx < len(table_data["rows"]):
                        run = row_cells[col_idx].paragraphs[0].runs[0]
                        run.bold = True
                        run.font.name = "Times New Roman"
                        run.font.size = Pt(10)
                    else:
                        run = row_cells[col_idx].paragraphs[0].runs[0]
                        run.bold = False
                        run.font.name = "Times New Roman"
                        run.font.size = Pt(10)

            # Gộp ô theo chiều ngang
            for merge_data in table_data.get("horizontal_merge", []):
                row = merge_data["row"]
                col_start = merge_data["col_start"]
                col_end = merge_data["col_end"]
                cell_to_keep = table.cell(row, col_start)

                # Xóa nội dung các ô bị merge
                for col_idx in range(col_start + 1, col_end + 1):
                    clear_cell_content(table.cell(row, col_idx))

                # Thực hiện merge
                cell_to_keep.merge(table.cell(row, col_end))

            # Gộp ô theo chiều dọc
            for merge_data in table_data.get("vertical_merge", []):
                col = merge_data["col"]
                row_start = merge_data["row_start"]
                row_end = merge_data["row_end"]
                cell_to_keep = table.cell(row_start, col)

                # Xóa nội dung các ô bị merge
                for row_idx in range(row_start + 1, row_end + 1):
                    clear_cell_content(table.cell(row_idx, col))

                # Thực hiện merge
                cell_to_keep.merge(table.cell(row_end, col))

    # Tạo file Word trong bộ nhớ
    word_io = BytesIO()
    doc.save(word_io)
    word_io.seek(0)
    return word_io

# API nhận JSON và trả về file Word
@app.route("/generate-report", methods=["POST"])
def generate_report():
    try:
        # Kiểm tra sự tồn tại của template
        if not os.path.exists(TEMPLATE_PATH):
            return {"error": "Template không tồn tại!"}, 404

        # Đọc template
        with open(TEMPLATE_PATH, "r", encoding="utf-8") as f:
            template = json.load(f)

        # Lấy dữ liệu từ client
        user_data = request.json

        # Tạo file Word từ template và dữ liệu người dùng
        word_file = create_word_file(template, user_data["data"])

        # Trả về file Word
        return send_file(word_file, as_attachment=True, download_name="bao_cao.docx")
    except Exception as e:
        return {"error": str(e)}, 500

# Chạy ứng dụng Flask
if __name__ == "__main__":
    app.run(debug=True)