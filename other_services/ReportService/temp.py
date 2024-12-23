import os
import json
from flask import Flask, request, send_file
from docx import Document
from io import BytesIO
from flask_cors import CORS
from docx.shared import Pt, Cm
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH

app = Flask(__name__)
CORS(app)

# Đường dẫn template JSON
TEMPLATE_PATH = "template/bao_cao_huan_luyen_thang.json"

# Hàm tạo file Word
def create_word_file(template, data):
    doc = Document()

    # Xoay trang ngang
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = section.page_height, section.page_width

    # Xử lý từng phần tử trong template
    for item in template["template"]:
        if item["type"] == "title":
            p = doc.add_paragraph(item["content"])
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run()
            run.bold = True
            run.font.name = "Times New Roman"
            run.font.size = Pt(14)
        elif item["type"] == "subtitle":
            p = doc.add_paragraph(item["content"])
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run()
            run.font.name = "Times New Roman"
            run.font.size = Pt(12)
        elif item["type"] == "table":
            table_data = item["data"]
            column_widths_cm = [Cm(24 * w / 100) for w in table_data["column_widths"]]
            rows = table_data["rows"] + data["rows"]

            # Tạo bảng
            table = doc.add_table(rows=len(rows), cols=len(rows[0]), style="Table Grid")

            # Cài đặt chiều rộng cột
            for col_idx, col_width in enumerate(column_widths_cm):
                for cell in table.columns[col_idx].cells:
                    cell.width = col_width

            # Điền dữ liệu vào bảng
            for row_idx, row in enumerate(rows):
                row_cells = table.rows[row_idx].cells
                for col_idx, cell_data in enumerate(row):
                    row_cells[col_idx].text = str(cell_data)

                    # In đậm hai dòng đầu tiên
                    if row_idx < len(table_data["rows"]):
                        run = row_cells[col_idx].paragraphs[0].runs[0]
                        run.bold = True
                        run.font.name = "Times New Roman"
                        run.font.size = Pt(10)

            # Gộp ô theo chiều ngang
            for merge_data in table_data.get("horizontal_merge", []):
                row = merge_data["row"]
                col_start = merge_data["col_start"]
                col_end = merge_data["col_end"]
                table.cell(row, col_start).merge(table.cell(row, col_end))

            # Gộp ô theo chiều dọc
            for merge_data in table_data.get("vertical_merge", []):
                col = merge_data["col"]
                row_start = merge_data["row_start"]
                row_end = merge_data["row_end"]
                table.cell(row_start, col).merge(table.cell(row_end, col))

    # Tạo file Word trong bộ nhớ
    word_io = BytesIO()
    doc.save(word_io)
    word_io.seek(0)
    return word_io

# API nhận JSON và trả về file Word
@app.route("/generate-report", methods=["POST"])
def generate_report():
    try:
        # Kiểm tra sự tồn tại của template
        if not os.path.exists(TEMPLATE_PATH):
            return {"error": "Template không tồn tại!"}, 404

        # Đọc template
        with open(TEMPLATE_PATH, "r", encoding="utf-8") as f:
            template = json.load(f)

        # Lấy dữ liệu từ client
        user_data = request.json

        # Tạo file Word từ template và dữ liệu người dùng
        word_file = create_word_file(template, user_data["data"])

        # Trả về file Word
        return send_file(word_file, as_attachment=True, download_name="bao_cao.docx")
    except Exception as e:
        return {"error": str(e)}, 500

# Chạy ứng dụng Flask
if __name__ == "__main__":
    app.run(debug=True)